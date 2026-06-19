"""Document parsing: LocalDocparse (pdf/docx/text, default) and DocAIDocparse (prod).

PDF/DOCX parsing uses the `docparse` extra (pypdf, python-docx). Plain text always works.
"""

from __future__ import annotations

import io

from captureos.config import Settings
from captureos.providers.base import DocparseProvider, ParsedDocument, ParsedPage


class LocalDocparse(DocparseProvider):
    name = "local"

    def __init__(self, settings: Settings) -> None:
        self._settings = settings

    async def parse(self, data: bytes, *, mime_type: str | None, filename: str) -> ParsedDocument:
        lname = filename.lower()
        mt = (mime_type or "").lower()

        if "pdf" in mt or lname.endswith(".pdf"):
            return self._parse_pdf(data)
        if "word" in mt or lname.endswith((".docx", ".doc")):
            return self._parse_docx(data)
        # Fallback: treat as UTF-8 text (also handles pasted solicitation text, FR-DI-3).
        text = data.decode("utf-8", errors="replace")
        return ParsedDocument(text=text, pages=[ParsedPage(page=1, text=text)], page_count=1)

    def _parse_pdf(self, data: bytes) -> ParsedDocument:
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("pypdf not installed (uv sync --extra docparse)") from exc
        reader = PdfReader(io.BytesIO(data))
        pages = [
            ParsedPage(page=i + 1, text=(p.extract_text() or ""))
            for i, p in enumerate(reader.pages)
        ]
        return ParsedDocument(
            text="\n\n".join(p.text for p in pages), pages=pages, page_count=len(pages)
        )

    def _parse_docx(self, data: bytes) -> ParsedDocument:
        try:
            import docx  # type: ignore
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("python-docx not installed (uv sync --extra docparse)") from exc
        document = docx.Document(io.BytesIO(data))
        text = "\n".join(p.text for p in document.paragraphs)
        return ParsedDocument(text=text, pages=[ParsedPage(page=1, text=text)], page_count=1)


class DocAIDocparse(DocparseProvider):  # pragma: no cover - requires GCP credentials
    name = "docai"

    def __init__(self, settings: Settings) -> None:
        if not settings.docai_processor_id:
            raise RuntimeError("DOCAI_PROCESSOR_ID required when DOCPARSE_PROVIDER=docai")
        self._settings = settings

    async def parse(self, data: bytes, *, mime_type: str | None, filename: str) -> ParsedDocument:
        import anyio
        from google.cloud import documentai  # type: ignore

        def _run() -> ParsedDocument:
            client = documentai.DocumentProcessorServiceClient()
            raw = documentai.RawDocument(content=data, mime_type=mime_type or "application/pdf")
            request = documentai.ProcessRequest(
                name=self._settings.docai_processor_id, raw_document=raw
            )
            result = client.process_document(request=request)
            doc = result.document
            pages = [ParsedPage(page=i + 1, text=doc.text) for i in range(len(doc.pages))]
            return ParsedDocument(text=doc.text, pages=pages, page_count=len(doc.pages))

        return await anyio.to_thread.run_sync(_run)
