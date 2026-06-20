"""Filing package export to Markdown / PDF / DOCX (FR-PB-5, CON-1).

Export only ever produces a downloadable artifact; it never transmits to any external system."""

from __future__ import annotations

import io
import textwrap

from docx import Document as DocxDocument
from fpdf import FPDF

MEDIA_TYPES = {
    "md": "text/markdown",
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

# Core PDF fonts are latin-1 only; transliterate the few non-latin1 glyphs we emit.
_LATIN1 = str.maketrans({"“": '"', "”": '"', "‘": "'", "’": "'", "–": "-", "—": "-", "•": "-"})


def _latinize(text: str) -> str:
    return text.translate(_LATIN1).encode("latin-1", "replace").decode("latin-1")


def _combined_markdown(title: str, docs: list[tuple[str, str]]) -> str:
    parts = [f"# {title}", ""]
    for _doc_type, content_md in docs:
        parts.append(content_md)
        parts.append("\n---\n")
    return "\n".join(parts)


def _emit(pdf: FPDF, text: str, *, size: int, bold: bool) -> None:
    # Wrap to a char width scaled to the font so even a long unbreakable token fits the page;
    # fpdf2 raises if a single glyph can't fit, so we never hand it an over-wide line.
    pdf.set_font("Helvetica", "B" if bold else "", size)
    width = max(24, int(880 / size))
    for chunk in textwrap.wrap(text, width=width, break_long_words=True) or [""]:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, size * 0.55 + 1, _latinize(chunk) or " ")


def _render_pdf(title: str, docs: list[tuple[str, str]]) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    _emit(pdf, title, size=16, bold=True)
    pdf.ln(2)
    for _doc_type, content_md in docs:
        for line in content_md.splitlines():
            if line.startswith("# "):
                _emit(pdf, line[2:].strip(), size=14, bold=True)
            elif line.startswith("## "):
                _emit(pdf, line[3:].strip(), size=12, bold=True)
            else:
                _emit(pdf, line, size=10, bold=False)
        pdf.ln(3)
    return bytes(pdf.output())  # bytearray in fpdf2 >= 2.8


def _render_docx(title: str, docs: list[tuple[str, str]]) -> bytes:
    document = DocxDocument()
    document.add_heading(title, level=0)
    for _doc_type, content_md in docs:
        for line in content_md.splitlines():
            if line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                document.add_heading(line[2:].strip(), level=1)
            elif line.strip().startswith(("- ", "* ")):
                document.add_paragraph(line.strip()[2:], style="List Bullet")
            elif line.strip():
                document.add_paragraph(line)
        document.add_page_break()
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_export(fmt: str, title: str, docs: list[tuple[str, str]]) -> tuple[bytes, str, str]:
    """Return (bytes, media_type, filename) for the requested format."""
    safe_title = "".join(c if c.isalnum() else "_" for c in title)[:48] or "filing_package"
    if fmt == "md":
        return (
            _combined_markdown(title, docs).encode("utf-8"),
            MEDIA_TYPES["md"],
            f"{safe_title}.md",
        )
    if fmt == "pdf":
        return _render_pdf(title, docs), MEDIA_TYPES["pdf"], f"{safe_title}.pdf"
    if fmt == "docx":
        return _render_docx(title, docs), MEDIA_TYPES["docx"], f"{safe_title}.docx"
    raise ValueError(f"Unsupported export format: {fmt}")
